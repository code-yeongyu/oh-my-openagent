#!/usr/bin/env python3
"""
Academic Review — Response-to-Authors DOCX Generator

Converts a markdown Response-to-Authors document into a professionally
formatted .docx file following academic journal standards.

Usage:
    python3 academic_review_docx.py <input.md> <output.docx>

Requires: python-docx (pip install python-docx)
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
        tcPr.append(tcBorders)
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = tcBorders.makeelement(
                qn(f"w:{edge}"),
                {
                    qn("w:val"): val.get("val", "single"),
                    qn("w:sz"): val.get("sz", "4"),
                    qn("w:color"): val.get("color", "000000"),
                    qn("w:space"): val.get("space", "0"),
                },
            )
            tcBorders.append(element)


def create_document():
    """Create a new document with academic formatting."""
    doc = Document()

    # Page margins — 1 inch all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # Paragraph spacing
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    # Heading styles
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

    return doc


def parse_markdown(md_text):
    """Parse markdown into structured blocks."""
    lines = md_text.split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,3})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            blocks.append({"type": "heading", "level": level, "text": m.group(2).strip()})
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i].lstrip("> ").strip())
                i += 1
            blocks.append({"type": "blockquote", "text": " ".join(quote_lines)})
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[-|\s:]+\|$", lines[i + 1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                # Skip separator line
                if re.match(r"^\|[-|\s:]+\|$", lines[i]):
                    i += 1
                    continue
                cells = [c.strip() for c in lines[i].split("|")[1:-1]]
                table_lines.append(cells)
                i += 1
            blocks.append({"type": "table", "rows": table_lines})
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "olist", "items": items})
            continue

        # Unordered list
        if re.match(r"^[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i]):
                items.append(re.sub(r"^[-*]\s+", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "ulist", "items": items})
            continue

        # Regular paragraph
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith(">") and not lines[i].startswith("|") and not re.match(r"^-{3,}$", lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        if para_lines:
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, code, links)."""
    # Pattern for **bold**, *italic*, `code`, [text](url)
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))"
    last_end = 0

    for m in re.finditer(pattern, text):
        # Add text before this match
        if m.start() > last_end:
            paragraph.add_run(text[last_end : m.start()])

        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):  # `code`
            run = paragraph.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif m.group(5):  # [text](url)
            run = paragraph.add_run(m.group(5))
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 238)

        last_end = m.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def build_docx(md_text, output_path, title="Response to Authors"):
    """Convert markdown to formatted .docx."""
    doc = create_document()
    blocks = parse_markdown(md_text)

    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = "Times New Roman"

    for block in blocks:
        btype = block["type"]

        if btype == "hr":
            # Add a thin line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(
                qn("w:bottom"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "6",
                    qn("w:space"): "1",
                    qn("w:color"): "999999",
                },
            )
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif btype == "heading":
            level = block["level"]
            p = doc.add_heading(level=level)
            add_formatted_text(p, block["text"])

        elif btype == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run(block["text"])
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)

        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"

            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < ncols:
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                        if ri == 0:
                            run.bold = True
                            # Shade header row
                            shading = cell._tc.get_or_add_tcPr().makeelement(
                                qn("w:shd"),
                                {
                                    qn("w:val"): "clear",
                                    qn("w:color"): "auto",
                                    qn("w:fill"): "E8E8E8",
                                },
                            )
                            cell._tc.get_or_add_tcPr().append(shading)

        elif btype == "olist":
            for idx, item in enumerate(block["items"], 1):
                p = doc.add_paragraph(style="List Number")
                add_formatted_text(p, item)

        elif btype == "ulist":
            for item in block["items"]:
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_text(p, item)

        elif btype == "paragraph":
            p = doc.add_paragraph()
            add_formatted_text(p, block["text"])

    # Add page numbers
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        fldChar1 = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._r.append(fldChar1)
        run2 = fp.add_run()
        instrText = run2._r.makeelement(qn("w:instrText"), {})
        instrText.text = " PAGE "
        run2._r.append(instrText)
        run3 = fp.add_run()
        fldChar2 = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3._r.append(fldChar2)

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 academic_review_docx.py <input.md> <output.docx>")
        print("       python3 academic_review_docx.py <input.md> <output.docx> --title 'Custom Title'")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    # Parse optional --title argument
    title = "Response to Authors"
    if "--title" in sys.argv:
        idx = sys.argv.index("--title")
        if idx + 1 < len(sys.argv):
            title = sys.argv[idx + 1]

    md_text = input_path.read_text(encoding="utf-8")
    result = build_docx(md_text, str(output_path), title=title)
    print(f"Generated: {result}")


if __name__ == "__main__":
    main()
